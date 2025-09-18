# y.py
# pip install python-docx
import argparse, os, sys, glob
from typing import Iterable, Union, Optional
from docx import Document
from docx.table import _Cell, Table

# 0-9 -> ০-৯
DIGIT_MAP = str.maketrans("0123456789", "০১২৩৪৫৬৭৮৯")

def translate_runs_in_paragraph(paragraph):
    for run in paragraph.runs:
        t = run.text
        if t and any(ch.isdigit() for ch in t):
            run.text = t.translate(DIGIT_MAP)

def iter_paragraphs_in_obj(obj: Union[Document, _Cell, Table]) -> Iterable:
    """Yield all paragraphs in Document/_Cell/Table (recursively handles tables)."""
    if hasattr(obj, "paragraphs"):
        for p in obj.paragraphs:
            yield p
    if hasattr(obj, "tables"):
        for tbl in obj.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    yield from iter_paragraphs_in_obj(cell)

def convert_docx_digits(src_path: str, dst_path: str) -> None:
    doc = Document(src_path)

    # Body + tables
    for p in iter_paragraphs_in_obj(doc):
        translate_runs_in_paragraph(p)

    # Headers & footers (per section)
    for section in doc.sections:
        for p in iter_paragraphs_in_obj(section.header):
            translate_runs_in_paragraph(p)
        for p in iter_paragraphs_in_obj(section.footer):
            translate_runs_in_paragraph(p)

    # Save
    outdir = os.path.dirname(dst_path)
    if outdir:
        os.makedirs(outdir, exist_ok=True)
    doc.save(dst_path)

def resolve_from_base(path: str, base: str) -> str:
    """Resolve `path` relative to script folder `base` if not absolute."""
    return path if os.path.isabs(path) else os.path.abspath(os.path.join(base, path))

def valid_docx(path: str) -> bool:
    return os.path.isfile(path) and path.lower().endswith(".docx")

def run_single(BASE: str, infile: str, outfile: Optional[str]):
    src = resolve_from_base(infile, BASE)
    if not valid_docx(src):
        print(f"[ERROR] Input file not found or not a .docx: {src}", file=sys.stderr)
        sys.exit(1)
    dst = resolve_from_base(outfile, BASE) if outfile else os.path.splitext(src)[0] + "_BN.docx"
    try:
        convert_docx_digits(src, dst)
        print(f"[OK] Saved: {dst}")
    except Exception as e:
        print(f"[ERROR] Failed to process '{src}': {e}", file=sys.stderr)
        sys.exit(2)

def run_batch(BASE: str, folder: str):
    indir = resolve_from_base(folder, BASE)
    if not os.path.isdir(indir):
        print(f"[ERROR] Folder not found: {indir}", file=sys.stderr)
        sys.exit(1)
    files = glob.glob(os.path.join(indir, "*.docx"))
    if not files:
        print(f"[INFO] No .docx files found in {indir}")
        return
    ok = fail = 0
    for src in files:
        if src.lower().endswith("_bn.docx"):
            continue
        dst = os.path.splitext(src)[0] + "_BN.docx"
        try:
            convert_docx_digits(src, dst)
            print(f"[OK] {os.path.basename(dst)}")
            ok += 1
        except Exception as e:
            print(f"[ERROR] {os.path.basename(src)} -> {e}", file=sys.stderr)
            fail += 1
    print(f"\nDone. Success: {ok}, Failed: {fail}")

def get_base_dir() -> str:
    # Works when run as a script or from notebooks/REPL where __file__ is missing
    try:
        return os.path.dirname(os.path.abspath(__file__))  # type: ignore[name-defined]
    except NameError:
        return os.path.abspath(os.getcwd())

def main():
    BASE = get_base_dir()

    ap = argparse.ArgumentParser(
        description="Convert English digits to Bangla digits (০-৯) in .docx files."
    )
    g = ap.add_mutually_exclusive_group()
    g.add_argument("--in", dest="infile", help="Input .docx file (relative to script or absolute)")
    g.add_argument("--dir", dest="indir", help="Process all .docx files in this folder")
    ap.add_argument("--out", dest="outfile", help="Output .docx (single-file mode only)")

    # If running in a notebook, argparse will try to parse Jupyter args; ignore them safely
    if hasattr(sys, "argv"):
        args = ap.parse_args(args=None if sys.argv[0].endswith(".py") else [])
    else:
        args = ap.parse_args([])

    if args.infile:
        run_single(BASE, args.infile, args.outfile)
        return

    if args.indir:
        run_batch(BASE, args.indir)
        return

    # Default: look for 'xy.docx' next to this script (or CWD in notebooks), write 'xy_BN.docx'
    default_in = "./xy.docx"
    default_out = "xy_BN.docx"
    print(f"[INFO] No arguments provided. Using default: {default_in} -> {default_out}")
    run_single(BASE, default_in, default_out)

if __name__ == "__main__":
    main()
